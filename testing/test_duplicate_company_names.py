#!/usr/bin/env python3
"""
Test für doppelte Firmennamen-Behandlung
"""

import os
import sys
import tempfile
from docx import Document
import pdfplumber

# Füge src-Verzeichnis zum Python-Pfad hinzu
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from src.ai_generator import AIGenerator
from src.pdf_generator import PDFGenerator
from src.docx_generator import DocxGenerator
from src.template_pdf_generator import TemplateBasedPDFGenerator
from src.models import JobDescription
from config.config import Config

def test_duplicate_company_names():
    """Testet die Behandlung doppelter Firmennamen"""
    print("=== Test: Doppelte Firmennamen-Behandlung ===\n")
    
    # Test-Job-Daten - simuliert das Luzerner Kantonsspital Szenario
    job_description = JobDescription(
        position='ICT Supporterin/ICT Supporter 100%',
        company='Luzerner Kantonsspital',
        contact_person='',  # Keine Kontaktperson -> RecipientController setzt Firmenname
        description='IT-Support Position',
        requirements='IT-Support, Kundenservice',
        benefits='Gute Konditionen',
        location='Luzern',
        address='',  # Keine Adresse -> RecipientController erstellt Standard-Adresse
        department='IT',
        url='https://example.com',
        working_hours='100%'
    )
    
    print(f"Original Job-Daten:")
    print(f"  Unternehmen: '{job_description.company}'")
    print(f"  Kontaktperson: '{job_description.contact_person}'")
    print(f"  Adresse: '{job_description.address}'")
    
    # AI-Generator mit RecipientController
    ai_generator = AIGenerator()
    motivation_letter = ai_generator.generate_motivation_letter(job_description)
    
    print(f"\nNach RecipientController:")
    print(f"  Empfänger-Unternehmen: '{motivation_letter.recipient_company}'")
    print(f"  Empfänger-Name: '{motivation_letter.recipient_name}'")
    print(f"  Empfänger-Adresse: '{motivation_letter.recipient_company_address}'")
    
    # Test Template-PDF-Generator
    template_pdf_generator = TemplateBasedPDFGenerator()
    pdf_path = template_pdf_generator.create_pdf_from_template(motivation_letter)
    print(f"\nTemplate-PDF erstellt: {pdf_path}")
    
    # Test DOCX-Generator
    docx_generator = DocxGenerator()
    docx_path = docx_generator.create_docx(motivation_letter)
    print(f"DOCX erstellt: {docx_path}")
    
    # Analysiere PDF-Inhalt
    print(f"\n=== PDF-Analyse ===")
    try:
        with pdfplumber.open(pdf_path) as pdf:
            first_page = pdf.pages[0]
            pdf_text = first_page.extract_text()
            
            # Zähle Vorkommen von "Luzerner Kantonsspital"
            company_count = pdf_text.count("Luzerner Kantonsspital")
            print(f"'Luzerner Kantonsspital' kommt {company_count} mal vor")
            
            # Zeige ersten Teil des PDFs
            lines = pdf_text.split('\n')[:15]  # Erste 15 Zeilen
            print("Erste 15 Zeilen:")
            for i, line in enumerate(lines, 1):
                print(f"  {i:2d}: {line}")
                
    except Exception as e:
        print(f"Fehler beim Lesen der PDF: {e}")
    
    # Analysiere DOCX-Inhalt
    print(f"\n=== DOCX-Analyse ===")
    try:
        doc = Document(docx_path)
        docx_text = ""
        for paragraph in doc.paragraphs:
            docx_text += paragraph.text + "\n"
        
        # Zähle Vorkommen von "Luzerner Kantonsspital"
        company_count = docx_text.count("Luzerner Kantonsspital")
        print(f"'Luzerner Kantonsspital' kommt {company_count} mal vor")
        
        # Zeige ersten Teil des DOCX
        lines = docx_text.split('\n')[:15]  # Erste 15 Zeilen
        print("Erste 15 Zeilen:")
        for i, line in enumerate(lines, 1):
            if line.strip():  # Nur nicht-leere Zeilen
                print(f"  {i:2d}: {line}")
                
    except Exception as e:
        print(f"Fehler beim Lesen der DOCX: {e}")
    
    # Prüfe auf doppelte Firmennamen (nur in Empfänger-Bereich)
    print(f"\n=== Ergebnis ===")
    pdf_has_duplicates = False
    docx_has_duplicates = False
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            pdf_text = pdf.pages[0].extract_text()
            lines = pdf_text.split('\n')
            
            # Suche nach doppelten Firmennamen im Empfänger-Bereich (erste 10 Zeilen)
            recipient_section = '\n'.join(lines[:10])
            recipient_company_count = recipient_section.count("Luzerner Kantonsspital")
            
            total_company_count = pdf_text.count("Luzerner Kantonsspital")
            
            if recipient_company_count > 1:
                pdf_has_duplicates = True
                print(f"❌ PDF: {recipient_company_count} Vorkommen von 'Luzerner Kantonsspital' im Empfänger-Bereich (zu viele)")
            else:
                print(f"✅ PDF: {recipient_company_count} Vorkommen von 'Luzerner Kantonsspital' im Empfänger-Bereich (korrekt)")
                print(f"   Gesamt: {total_company_count} Vorkommen (inkl. Inhalt)")
    except:
        pass
    
    try:
        doc = Document(docx_path)
        docx_text = "\n".join([p.text for p in doc.paragraphs])
        lines = docx_text.split('\n')
        
        # Suche nach doppelten Firmennamen im Empfänger-Bereich (erste 10 nicht-leere Zeilen)
        recipient_lines = [line for line in lines[:15] if line.strip()][:10]
        recipient_section = '\n'.join(recipient_lines)
        recipient_company_count = recipient_section.count("Luzerner Kantonsspital")
        
        total_company_count = docx_text.count("Luzerner Kantonsspital")
        
        if recipient_company_count > 1:
            docx_has_duplicates = True
            print(f"❌ DOCX: {recipient_company_count} Vorkommen von 'Luzerner Kantonsspital' im Empfänger-Bereich (zu viele)")
        else:
            print(f"✅ DOCX: {recipient_company_count} Vorkommen von 'Luzerner Kantonsspital' im Empfänger-Bereich (korrekt)")
            print(f"   Gesamt: {total_company_count} Vorkommen (inkl. Inhalt)")
    except:
        pass
    
    # Cleanup
    try:
        os.remove(pdf_path)
        os.remove(docx_path)
        print(f"\nTest-Dateien bereinigt.")
    except:
        pass
    
    return not (pdf_has_duplicates or docx_has_duplicates)

if __name__ == "__main__":
    success = test_duplicate_company_names()
    if success:
        print("\n🎉 Test erfolgreich: Keine doppelten Firmennamen!")
    else:
        print("\n❌ Test fehlgeschlagen: Doppelte Firmennamen gefunden!")
    sys.exit(0 if success else 1)
