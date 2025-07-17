#!/usr/bin/env python3
"""
Test für Hyperlink-Funktionalität in DOCX-Dokumenten
"""

from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
import os

def test_hyperlink_creation():
    """Testet die Hyperlink-Erstellung"""
    doc = Document()
    
    # Einfacher Hyperlink-Test
    paragraph = doc.add_paragraph()
    
    # Methode 1: Manueller Hyperlink
    try:
        # Neue Hyperlink-Relation erstellen
        part = paragraph.part
        r_id = part.relate_to("https://github.com/zurd46/ZurdLLMWS", 
                             "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", 
                             is_external=True)
        
        print(f"Relation ID erstellt: {r_id}")
        
        # Hyperlink-XML erstellen
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # Run-Element für den Hyperlink-Text
        run = OxmlElement('w:r')
        
        # Run-Properties für blaue Farbe und Unterstreichung
        run_props = OxmlElement('w:rPr')
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0000FF')  # Blau
        underline = OxmlElement('w:u')
        underline.set(qn('w:val'), 'single')
        
        run_props.append(color)
        run_props.append(underline)
        run.append(run_props)
        
        # Text-Element
        text_elem = OxmlElement('w:t')
        text_elem.text = "ZurdLLMWS"
        run.append(text_elem)
        
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
        
        print("Hyperlink erfolgreich erstellt!")
        
    except Exception as e:
        print(f"Fehler bei Hyperlink-Erstellung: {e}")
        # Fallback: Normaler Text
        run = paragraph.add_run("ZurdLLMWS (Fallback)")
        run.font.size = Pt(11)
        run.font.color.rgb = None
    
    # Normaler Text danach
    paragraph.add_run(" - Ein Test-Hyperlink")
    
    # Speichern
    output_path = "output/hyperlink_test.docx"
    doc.save(output_path)
    print(f"Test-Dokument gespeichert: {output_path}")
    
    return output_path

def analyze_hyperlink_structure(docx_path):
    """Analysiert die Struktur des DOCX-Dokuments"""
    import zipfile
    
    if not os.path.exists(docx_path):
        print(f"Datei nicht gefunden: {docx_path}")
        return
    
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # Relationships prüfen
        try:
            rels_xml = zip_ref.read('word/_rels/document.xml.rels')
            print("\n=== Relationships ===")
            print(rels_xml.decode('utf-8'))
        except Exception as e:
            print(f"Fehler beim Lesen der Relationships: {e}")
        
        # Document XML prüfen
        try:
            doc_xml = zip_ref.read('word/document.xml')
            content = doc_xml.decode('utf-8')
            print("\n=== Document XML (erste 1000 Zeichen) ===")
            print(content[:1000])
            
            # Nach Hyperlinks suchen
            import re
            hyperlinks = re.findall(r'<w:hyperlink[^>]*>.*?</w:hyperlink>', content, re.DOTALL)
            print(f"\nAnzahl Hyperlinks gefunden: {len(hyperlinks)}")
            for i, link in enumerate(hyperlinks):
                print(f"Hyperlink {i+1}: {link}")
        except Exception as e:
            print(f"Fehler beim Lesen des Document XML: {e}")

if __name__ == "__main__":
    print("=== Hyperlink-Test ===")
    test_path = test_hyperlink_creation()
    analyze_hyperlink_structure(test_path)
