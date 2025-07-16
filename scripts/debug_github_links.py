#!/usr/bin/env python3
"""
Debug-Skript für GitHub-Links in DOCX
"""

from src.docx_generator import DocxGenerator
from src.models import MotivationLetter

# Test-Motivationsschreiben mit GitHub-Projekten
test_letter = MotivationLetter(
    sender_name='Daniel Zurmühle',
    sender_address='Hinterdorfstrasse 12, 6235 Winikon',
    sender_phone='+41 79 127 55 54',
    sender_email='dzurmuehle@gmail.com',
    recipient_company='GitHub Test Company',
    recipient_company_address='Test Address',
    recipient_address='Test Address',
    recipient_name='Test Person',
    subject='Bewerbung als Softwareentwickler',
    content='Sehr geehrte Damen und Herren,\n\nIch habe bereits an mehreren Projekten gearbeitet, darunter AutomaticMotivation und ZurdLLMWS. Diese Projekte zeigen meine Fähigkeiten in der Softwareentwicklung.\n\nMein LinkedIn-Profil finden Sie unter meinem Namen.\n\nMit freundlichen Grüßen'
)

# Debug der Paragraph-Verarbeitung
print("Debug: Content-Verarbeitung")
print(f"Original content: {test_letter.content}")
print("\nParagraph-Aufteilung:")
paragraphs = test_letter.content.split('\n\n')
for i, paragraph_text in enumerate(paragraphs):
    if paragraph_text.strip():
        cleaned_text = paragraph_text.strip()
        print(f"Paragraph {i}: '{cleaned_text}'")
        
        # Check greeting filter
        if (cleaned_text.startswith('Sehr geehrte Damen und Herren') or
            cleaned_text.startswith('Sehr geehrter Herr') or
            cleaned_text.startswith('Sehr geehrte Frau') or
            cleaned_text in ['Sehr geehrte Damen und Herren,', 'Sehr geehrte Damen und Herren']):
            print(f"  -> FILTERED OUT (greeting)")
            continue
        
        # Check closing filter
        if (cleaned_text.startswith('Mit freundlichen Grüßen') or
            cleaned_text.startswith('Mit freundlichen Grüssen') or
            cleaned_text.startswith('Freundliche Grüße') or
            cleaned_text.startswith('Freundliche Grüsse') or
            cleaned_text in ['Mit freundlichen Grüßen', 'Mit freundlichen Grüssen']):
            print(f"  -> FILTERED OUT (closing)")
            continue
        
        print(f"  -> WOULD BE ADDED TO DOCX")

print("\nGeneriere DOCX...")
generator = DocxGenerator()
filepath = generator.create_docx(test_letter)
print(f"DOCX erstellt: {filepath}")

# Prüfe das erstellte DOCX
from docx import Document
doc = Document(filepath)
print(f"\nDOCX Inhalt ({len(doc.paragraphs)} Paragraphen):")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"{i}: {text}")
        if 'hyperlink' in str(para._element.xml):
            print(f"  -> Hyperlink found in paragraph {i}")
