#!/usr/bin/env python3
"""
DOCX Generator für AutomaticMotivation
Erstellt Microsoft Word-Dokumente basierend auf Motivationsschreiben
"""

import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from src.models import MotivationLetter
from config.config import Config
import re
import logging

logger = logging.getLogger(__name__)

class DocxGenerator:
    """Klasse zur Erstellung von DOCX-Dokumenten für Motivationsschreiben"""
    
    def __init__(self):
        """Initialisiert den DOCX Generator"""
        self.output_dir = "output"
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
        
        # Konfiguriere bevorzugte Schriftart
        self.font_name = 'Aptos Display'
        self.font_fallback = 'Calibri'  # Fallback für ältere Word-Versionen
    
    def _set_font(self, run, size: int = 11, bold: bool = False):
        """Setzt die Schriftart für einen Run"""
        run.font.size = Pt(size)
        run.font.bold = bold
        try:
            run.font.name = self.font_name
        except:
            # Fallback zu Calibri wenn Aptos Display nicht verfügbar
            run.font.name = self.font_fallback
    
    def create_docx(self, motivation_letter: MotivationLetter) -> str:
        """
        Erstellt ein DOCX-Dokument aus einem Motivationsschreiben
        
        Args:
            motivation_letter: Das Motivationsschreiben-Objekt
            
        Returns:
            str: Pfad zur erstellten DOCX-Datei
        """
        try:
            # Neues Dokument erstellen
            doc = Document()
            
            # Seitenränder setzen (2.5cm wie bei PDF)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            # Absender-Informationen (oben rechts)
            self._add_sender_info(doc, motivation_letter)
            
            # Leerzeile
            doc.add_paragraph()
            
            # Empfänger-Informationen
            self._add_recipient_info(doc, motivation_letter)
            
            # Leerzeile
            doc.add_paragraph()
            
            # Datum und Ort
            self._add_date_location(doc, motivation_letter)
            
            # Leerzeile
            doc.add_paragraph()
            
            # Betreff
            self._add_subject(doc, motivation_letter)
            
            # Leerzeile
            doc.add_paragraph()
            
            # Anrede
            self._add_salutation(doc, motivation_letter)
            
            # Hauptinhalt
            self._add_main_content(doc, motivation_letter)
            
            # Grußformel
            self._add_closing(doc, motivation_letter)
            
            # Dateiname generieren
            filename = self._generate_filename(motivation_letter)
            filepath = os.path.join(self.output_dir, filename)
            
            # Dokument speichern
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Fehler beim Erstellen der DOCX-Datei: {e}")
    
    def _add_sender_info(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Absender-Informationen hinzu"""
        # Absender-Informationen rechtsbündig
        sender_paragraph = doc.add_paragraph()
        sender_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Kompakte Formatierung für Absender
        sender_format = sender_paragraph.paragraph_format
        sender_format.space_after = Pt(0)
        sender_format.line_spacing = 1.0
        
        # Name
        sender_run = sender_paragraph.add_run(f"{motivation_letter.sender_name}\n")
        self._set_font(sender_run, 11)
        
        # Adresse in separate Zeilen aufteilen
        address_parts = motivation_letter.sender_address.split(', ')
        
        if len(address_parts) >= 2:
            # Straße und Hausnummer (erster Teil)
            street = address_parts[0].strip()
            address_run = sender_paragraph.add_run(f"{street}\n")
            self._set_font(address_run, 11)
            
            # PLZ und Stadt (zweiter Teil)
            city_part = address_parts[1].strip()
            city_run = sender_paragraph.add_run(f"{city_part}\n")
            self._set_font(city_run, 11)
        else:
            # Fallback: Komplette Adresse in einer Zeile
            address_run = sender_paragraph.add_run(f"{motivation_letter.sender_address}\n")
            self._set_font(address_run, 11)
        
        # Telefonnummer
        contact_run = sender_paragraph.add_run(f"{motivation_letter.sender_phone}\n")
        contact_run.font.size = Pt(11)
        contact_run.font.name = 'Aptos Display'
        
        # E-Mail
        email_run = sender_paragraph.add_run(f"{motivation_letter.sender_email}")
        email_run.font.size = Pt(11)
        email_run.font.name = 'Aptos Display'
    
    def _add_recipient_info(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Empfänger-Informationen hinzu"""
        # Firmenname (normal, nicht fett und nicht unterstrichen)
        company_paragraph = doc.add_paragraph()
        company_run = company_paragraph.add_run(f"{motivation_letter.recipient_company}")
        company_run.font.size = Pt(11)
        company_run.font.name = 'Aptos Display'
        company_run.bold = False  # Nicht fett
        company_run.underline = False  # Nicht unterstrichen
        
        # Zeilenabstand für Firmenname reduzieren
        company_format = company_paragraph.paragraph_format
        company_format.space_after = Pt(0)
        company_format.line_spacing = 1.0
        
        # Adresse in separaten Zeilen formatieren
        if motivation_letter.recipient_company_address and motivation_letter.recipient_company_address != "Nicht angegeben":
            # Adresse nach Kommas aufteilen und jede Zeile separat hinzufügen
            address_parts = motivation_letter.recipient_company_address.split(',')
            
            for part in address_parts:
                part = part.strip()
                if part:
                    address_paragraph = doc.add_paragraph()
                    address_run = address_paragraph.add_run(part)
                    address_run.font.size = Pt(11)
                    address_run.font.name = 'Aptos Display'
                    
                    # Zeilenabstand für Adresse reduzieren
                    address_format = address_paragraph.paragraph_format
                    address_format.space_after = Pt(0)
                    address_format.space_before = Pt(0)
                    address_format.line_spacing = 1.0
    
    def _add_date_location(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Datum und Ort hinzu"""
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Ort aus Adresse extrahieren
        location = self._extract_location_from_address(motivation_letter.sender_address)
        current_date = datetime.now().strftime("%d.%m.%Y")
        
        date_run = date_paragraph.add_run(f"{location}, {current_date}")
        date_run.font.size = Pt(11)
        date_run.font.name = 'Aptos Display'
    
    def _add_subject(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Betreff hinzu"""
        subject_paragraph = doc.add_paragraph()
        subject_run = subject_paragraph.add_run(f"{motivation_letter.subject}")
        subject_run.font.size = Pt(11)
        subject_run.font.name = 'Aptos Display'
        subject_run.bold = True
    
    def _add_salutation(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Anrede hinzu"""
        salutation_paragraph = doc.add_paragraph()
        
        # Verbesserte Anrede-Logik
        salutation = "Sehr geehrte Damen und Herren,"
        
        if motivation_letter.recipient_name and motivation_letter.recipient_name != motivation_letter.recipient_company:
            recipient_name = motivation_letter.recipient_name.strip()
            recipient_lower = recipient_name.lower()
            
            # Überprüfe auf männliche Anrede - erweiterte Logik
            if (recipient_lower.startswith('herr ') or 
                recipient_lower.startswith('hr. ') or 
                recipient_lower.startswith('mr. ') or
                ' herr ' in recipient_lower or
                recipient_lower.startswith('jan ') or  # Für Jan am Anfang
                'jan ' in recipient_lower or  # Für Jan irgendwo
                any(name in recipient_lower for name in ['schmitz-elsen', 'jan schmitz-elsen', 'schmitz', 'elsen'])):
                
                # Entferne Titel für saubere Anrede
                clean_name = recipient_name
                if recipient_name.lower().startswith('herr '):
                    clean_name = recipient_name[5:]
                elif recipient_name.lower().startswith('hr. '):
                    clean_name = recipient_name[4:]
                elif recipient_name.lower().startswith('mr. '):
                    clean_name = recipient_name[4:]
                
                # Für Jan Schmitz-Elsen: Nur Nachname verwenden
                if 'jan' in recipient_lower and 'schmitz-elsen' in recipient_lower:
                    clean_name = "Schmitz-Elsen"
                elif 'jan' in recipient_lower and 'schmitz' in recipient_lower:
                    clean_name = "Schmitz-Elsen"
                
                salutation = f"Sehr geehrter Herr {clean_name},"
            
            # Überprüfe auf weibliche Anrede
            elif (recipient_lower.startswith('frau ') or 
                  recipient_lower.startswith('fr. ') or 
                  recipient_lower.startswith('mrs. ') or
                  recipient_lower.startswith('ms. ') or
                  ' frau ' in recipient_lower):
                
                # Entferne Titel für saubere Anrede
                clean_name = recipient_name
                if recipient_name.lower().startswith('frau '):
                    clean_name = recipient_name[5:]
                elif recipient_name.lower().startswith('fr. '):
                    clean_name = recipient_name[4:]
                elif recipient_name.lower().startswith('mrs. '):
                    clean_name = recipient_name[5:]
                elif recipient_name.lower().startswith('ms. '):
                    clean_name = recipient_name[4:]
                
                salutation = f"Sehr geehrte Frau {clean_name},"
        
        salutation_run = salutation_paragraph.add_run(salutation)
        salutation_run.font.size = Pt(11)
        salutation_run.font.name = 'Aptos Display'
    
    def _add_main_content(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Hauptinhalt hinzu mit GitHub-Projekt-Hyperlinks"""
        # Motivationsschreiben in Absätze aufteilen
        paragraphs = motivation_letter.content.split('\n\n')
        
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                # Entferne doppelte Anreden aus dem Inhalt
                cleaned_text = paragraph_text.strip()
                
                # Überspringe Absätze, die nur Anreden enthalten
                if (cleaned_text.startswith('Sehr geehrte Damen und Herren') or
                    cleaned_text.startswith('Sehr geehrter Herr') or
                    cleaned_text.startswith('Sehr geehrte Frau') or
                    cleaned_text in ['Sehr geehrte Damen und Herren,', 'Sehr geehrte Damen und Herren']):
                    continue
                
                # Überspringe Grußformeln, da sie separat hinzugefügt werden
                if (cleaned_text.startswith('Mit freundlichen Grüßen') or
                    cleaned_text.startswith('Mit freundlichen Grüssen') or
                    cleaned_text.startswith('Freundliche Grüße') or
                    cleaned_text.startswith('Freundliche Grüsse') or
                    cleaned_text in ['Mit freundlichen Grüßen', 'Mit freundlichen Grüssen']):
                    continue
                
                content_paragraph = doc.add_paragraph()
                
                # Prüfe auf GitHub-Projekt-Erwähnungen und erstelle Hyperlinks
                self._add_text_with_github_links(content_paragraph, cleaned_text)
                
                # Zeilenabstand setzen
                paragraph_format = content_paragraph.paragraph_format
                paragraph_format.space_after = Pt(6)
                paragraph_format.line_spacing = 1.15
    
    def _add_text_with_github_links(self, paragraph, text):
        """Fügt Text mit GitHub-Projekt-Hyperlinks und LinkedIn-Links hinzu"""
        # Dynamische GitHub-Projekt-URLs aus Config laden
        project_urls = Config.get_github_project_urls()
        
        # Erstelle Regex-Pattern dynamisch aus verfügbaren Projekten
        if project_urls:
            # Escape spezielle Regex-Zeichen in Projektnamen
            escaped_projects = [re.escape(project) for project in project_urls.keys()]
            github_project_pattern = f"({'|'.join(escaped_projects)})"
        else:
            # Fallback-Pattern falls keine Projekte verfügbar
            github_project_pattern = r'(AutomaticMotivation|ZurdLLMWS|Auto-search-jobs)'
        
        # Regex-Pattern für LinkedIn-Erwähnungen
        # Sucht nach: "LinkedIn-Profil" oder "LinkedIn Profil"
        linkedin_pattern = r"(LinkedIn[\s\-]?Profil)"
        
        # Verarbeite zuerst GitHub-Projekte, dann LinkedIn
        processed_text = text
        
        # GitHub-Projekt-Links verarbeiten
        def replace_project_with_link(match):
            project_name = match.group(1)    # "ProjectName"
            
            # GitHub-URL für das Projekt suchen
            github_url = self._get_github_url_for_project(project_name)
            
            if github_url:
                return f'GITHUB_LINK:{project_name}:{github_url}'
            else:
                return project_name
        
        # LinkedIn-Links verarbeiten
        def replace_linkedin_with_link(match):
            linkedin_text = match.group(1)
            linkedin_url = Config.PERSONAL_LINKEDIN
            if linkedin_url:
                return f'LINKEDIN_LINK:{linkedin_text}:{linkedin_url}'
            else:
                return linkedin_text
        
        # Ersetze GitHub-Projekte
        processed_text = re.sub(github_project_pattern, replace_project_with_link, processed_text)
        
        # Ersetze LinkedIn-Erwähnungen
        processed_text = re.sub(linkedin_pattern, replace_linkedin_with_link, processed_text)
        
        # Jetzt den Text mit den Markierungen verarbeiten
        self._process_text_with_links(paragraph, processed_text)
    
    def _process_text_with_links(self, paragraph, text):
        """Verarbeitet Text mit GitHub- und LinkedIn-Link-Markierungen"""
        # Split text by link markers - berücksichtigt URLs mit ://
        parts = re.split(r'(GITHUB_LINK:[^:]+:https?://[^\s]+|LINKEDIN_LINK:[^:]+:https?://[^\s]+)', text)
        
        for part in parts:
            if part.startswith('GITHUB_LINK:'):
                # Extract GitHub link info
                try:
                    _, link_text, url = part.split(':', 2)
                    self._add_hyperlink(paragraph, link_text, url)
                except ValueError:
                    # Fallback bei Parsing-Fehlern
                    run = paragraph.add_run(part)
                    self._set_font(run, size=11, bold=False)
            elif part.startswith('LINKEDIN_LINK:'):
                # Extract LinkedIn link info
                try:
                    _, link_text, url = part.split(':', 2)
                    self._add_hyperlink(paragraph, link_text, url)
                except ValueError:
                    # Fallback bei Parsing-Fehlern
                    run = paragraph.add_run(part)
                    self._set_font(run, size=11, bold=False)
            else:
                # Regular text
                if part:
                    run = paragraph.add_run(part)
                    self._set_font(run, size=11, bold=False)
    
    def _extract_github_projects_from_text(self, text):
        """Extrahiert GitHub-Projektnamen aus dem Text"""
        # Diese Methode könnte erweitert werden, um dynamisch Projekte zu erkennen
        common_projects = ['AutomaticMotivation', 'ZurdLLMWS', 'Auto-search-jobs']
        found_projects = []
        
        for project in common_projects:
            if project in text:
                found_projects.append(project)
        
        return found_projects
    
    def _get_github_url_for_project(self, project_name):
        """Gibt die GitHub-URL für ein Projekt zurück"""
        # Verwende Config für dynamische URL-Generierung
        project_urls = Config.get_github_project_urls()
        return project_urls.get(project_name)
    
    def _add_hyperlink(self, paragraph, text, url):
        """Fügt einen Hyperlink zu einem Absatz hinzu"""
        try:
            # Validiere die URL
            if not url or not url.startswith(('http://', 'https://')):
                # Fallback: Normaler Text wenn URL ungültig
                run = paragraph.add_run(text)
                self._set_font(run, size=11, bold=False)
                return
            
            # Neue Hyperlink-Relation erstellen
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            
            # Hyperlink-XML erstellen
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Run-Element für den Hyperlink-Text
            run = OxmlElement('w:r')
            
            # Run-Properties für blaue Farbe und Unterstreichung
            run_props = OxmlElement('w:rPr')
            
            # Font-Familie
            font_family = OxmlElement('w:rFonts')
            font_family.set(qn('w:ascii'), 'Aptos Display')
            font_family.set(qn('w:hAnsi'), 'Aptos Display')
            run_props.append(font_family)
            
            # Schriftgröße
            font_size = OxmlElement('w:sz')
            font_size.set(qn('w:val'), '22')  # 11pt * 2
            run_props.append(font_size)
            
            # Blaue Farbe
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Blau
            run_props.append(color)
            
            # Unterstreichung
            underline = OxmlElement('w:u')
            underline.set(qn('w:val'), 'single')
            run_props.append(underline)
            
            run.append(run_props)
            
            # Text-Element
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            run.append(text_elem)
            
            hyperlink.append(run)
            paragraph._p.append(hyperlink)
            
            logger.info(f"Hyperlink erstellt: {text} -> {url}")
            
        except Exception as e:
            logger.error(f"Fehler bei Hyperlink-Erstellung für {text}: {e}")
            # Fallback: Normaler Text wenn Hyperlink fehlschlägt
            run = paragraph.add_run(text)
            self._set_font(run, size=11, bold=False)
    
    def _add_closing(self, doc: Document, motivation_letter: MotivationLetter):
        """Fügt Grußformel hinzu"""
        # Leerzeile vor Grußformel
        doc.add_paragraph()
        
        # Grußformel
        closing_paragraph = doc.add_paragraph()
        closing_run = closing_paragraph.add_run("Mit freundlichen Grüßen")
        closing_run.font.size = Pt(11)
        closing_run.font.name = 'Aptos Display'
        
        # Eine Leerzeile für Unterschrift (reduzierter Abstand)
        doc.add_paragraph()
        
        # Name für Unterschrift
        signature_paragraph = doc.add_paragraph()
        signature_run = signature_paragraph.add_run(f"{motivation_letter.sender_name}")
        signature_run.font.size = Pt(11)
        signature_run.font.name = 'Aptos Display'
    
    def _extract_location_from_address(self, address: str) -> str:
        """Extrahiert den Ort aus der Adresse"""
        if not address:
            return "Ort"
        
        # Verschiedene Muster für deutsche Adressen
        patterns = [
            r'\d{4,5}\s+([A-ZÄÖÜ][a-zäöüß-]+(?:\s+[A-ZÄÖÜ][a-zäöüß-]+)*)',  # PLZ + Stadt
            r',\s*([A-ZÄÖÜ][a-zäöüß-]+(?:\s+[A-ZÄÖÜ][a-zäöüß-]+)*)(?:\s*,|\s*$)',  # Nach Komma
        ]
        
        for pattern in patterns:
            match = re.search(pattern, address)
            if match:
                return match.group(1).strip()
        
        # Fallback: Letztes Wort nach Komma oder Leerzeichen
        parts = address.replace(',', ' ').split()
        if parts:
            return parts[-1]
        
        return "Ort"
    
    def _generate_filename(self, motivation_letter: MotivationLetter) -> str:
        """Generiert den Dateinamen für die DOCX-Datei"""
        # Firmenname bereinigen
        company_name = motivation_letter.recipient_company.replace(" ", "_")
        company_name = re.sub(r'[^\w\-_\.]', '', company_name)
        
        # Ort extrahieren
        location = self._extract_location_from_company_address(motivation_letter.recipient_company_address)
        if location:
            location = location.replace(" ", "_")
            location = re.sub(r'[^\w\-_\.]', '', location)
        
        # Datum formatieren OHNE Zeitstempel (nur Datum)
        date_str = datetime.now().strftime("%d%m%y")
        
        # Dateiname zusammensetzen
        if location:
            filename = f"Motivationsschreiben_{company_name}_{location}_{date_str}.docx"
        else:
            filename = f"Motivationsschreiben_{company_name}_{date_str}.docx"
        
        return filename
    
    def _extract_location_from_company_address(self, address: str) -> Optional[str]:
        """Extrahiert den Ort aus der Firmenadresse"""
        if not address or address == "Nicht angegeben":
            return None
        
        # Deutsche PLZ + Stadt Muster
        plz_stadt_pattern = r'\d{4,5}\s+([A-ZÄÖÜ][a-zäöüß-]+(?:\s+[A-ZÄÖÜ][a-zäöüß-]+)*)'
        match = re.search(plz_stadt_pattern, address)
        if match:
            return match.group(1).strip()
        
        # Schweizer PLZ + Stadt Muster
        ch_pattern = r'(\d{4})\s+([A-ZÄÖÜ][a-zäöüß-]+(?:\s+[A-ZÄÖÜ][a-zäöüß-]+)*)'
        match = re.search(ch_pattern, address)
        if match:
            return match.group(2).strip()
        
        return None
